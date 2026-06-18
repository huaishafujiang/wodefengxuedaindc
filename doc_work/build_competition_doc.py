from __future__ import annotations

import json
import math
import re
from pathlib import Path

import pypdfium2 as pdfium
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"D:\ai\doc_work")
INPUTS = ROOT / "inputs"
ASSETS = ROOT / "assets"
OUT = ROOT / "output"
TEMPLATE = INPUTS / "template.docx"
DOCX_OUT = OUT / "STM32G431智能频响测量与AI辅助电路诊断系统_要求文档.docx"

FONT_CN = "宋体"
FONT_EN = "Times New Roman"
FONT_CN_UI = r"C:\Windows\Fonts\msyh.ttc"
FONT_CN_UI_BOLD = r"C:\Windows\Fonts\msyhbd.ttc"

PAGE_USABLE_WIDTH_CM = 21.0 - 3.17 - 3.17


def ensure_dirs() -> None:
    ASSETS.mkdir(parents=True, exist_ok=True)
    OUT.mkdir(parents=True, exist_ok=True)


def set_run_font(run, size_pt=12, bold=False, color=None, cn_font=FONT_CN, en_font=FONT_EN) -> None:
    run.font.name = en_font
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), en_font)
    rfonts.set(qn("w:hAnsi"), en_font)
    rfonts.set(qn("w:cs"), en_font)
    rfonts.set(qn("w:eastAsia"), cn_font)
    rfonts.set(qn("w:hint"), "eastAsia")


def set_para_format(
    para,
    align=None,
    first_line_cm=None,
    line_spacing=1.5,
    before_pt=0,
    after_pt=0,
) -> None:
    para.alignment = align
    fmt = para.paragraph_format
    fmt.line_spacing = line_spacing
    fmt.space_before = Pt(before_pt)
    fmt.space_after = Pt(after_pt)
    if first_line_cm is not None:
        fmt.first_line_indent = Cm(first_line_cm)
    else:
        fmt.first_line_indent = None


def add_text_para(doc, text, *, style="Normal", size=12, bold=False, align=None, first_line=True):
    para = doc.add_paragraph(style=style)
    set_para_format(para, align=align, first_line_cm=0.85 if first_line else None)
    run = para.add_run(text)
    set_run_font(run, size_pt=size, bold=bold)
    return para


def add_section_heading(doc, prefix, title):
    para = doc.add_paragraph(style="Normal")
    set_para_format(para, first_line_cm=None)
    r1 = para.add_run(prefix + "  ")
    set_run_font(r1, size_pt=14, bold=True)
    r2 = para.add_run(title)
    set_run_font(r2, size_pt=12, bold=True)
    return para


def add_subheading(doc, text):
    para = doc.add_paragraph(style="List Paragraph")
    set_para_format(para, first_line_cm=None)
    run = para.add_run(text)
    set_run_font(run, size_pt=12, bold=False)
    return para


def add_caption(doc, text):
    para = doc.add_paragraph(style="Normal")
    set_para_format(para, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_cm=None)
    run = para.add_run(text)
    set_run_font(run, size_pt=10.5, bold=False)
    return para


def add_figure(doc, image_path: Path, caption: str, width_cm: float = PAGE_USABLE_WIDTH_CM):
    para = doc.add_paragraph(style="Normal")
    set_para_format(para, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_cm=None, before_pt=0, after_pt=0)
    run = para.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    add_caption(doc, caption)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=95, bottom=70, end=95) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for para in cell.paragraphs:
                set_para_format(para, first_line_cm=None, line_spacing=1.2)
                for run in para.runs:
                    set_run_font(run, size_pt=10.5)
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(sum(widths_cm) / 2.54 * 1440)))
    tbl_grid = tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width / 2.54 * 1440)))
        tbl_grid.append(grid_col)


def add_table(doc, headers, rows, widths_cm):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        table.style = "Normal Table"
    apply_table_borders(table)
    hdr = table.rows[0].cells
    for i, head in enumerate(headers):
        hdr[i].text = head
        set_cell_shading(hdr[i], "EDEDED")
        for para in hdr[i].paragraphs:
            set_para_format(para, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_cm=None, line_spacing=1.2)
            for run in para.runs:
                set_run_font(run, size_pt=10.5, bold=True)
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = str(val)
            for para in cells[i].paragraphs:
                set_para_format(
                    para,
                    align=WD_ALIGN_PARAGRAPH.CENTER if len(str(val)) <= 18 else None,
                    first_line_cm=None,
                    line_spacing=1.2,
                )
                for run in para.runs:
                    set_run_font(run, size_pt=10.5)
    set_table_widths(table, widths_cm)
    doc.add_paragraph()
    return table


def apply_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "7F7F7F")


def clear_document_body(doc: Document) -> None:
    body = doc._body._element
    sect_pr = None
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            sect_pr = child
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


def configure_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)
    sec.header_distance = Cm(1.5)
    sec.footer_distance = Cm(1.75)
    for style_name in ["Normal", "List Paragraph"]:
        style = doc.styles[style_name]
        style.font.name = FONT_EN
        style.font.size = Pt(12)
        rpr = style._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), FONT_EN)
        rfonts.set(qn("w:hAnsi"), FONT_EN)
        rfonts.set(qn("w:eastAsia"), FONT_CN)
        pf = style.paragraph_format
        pf.line_spacing = 1.5
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)


def render_pdf_page(pdf_path: Path, page_index: int, out_path: Path, scale: float = 4.0) -> None:
    pdf = pdfium.PdfDocument(str(pdf_path))
    page = pdf[page_index]
    bitmap = page.render(scale=scale).to_pil()
    if bitmap.mode != "RGB":
        bitmap = bitmap.convert("RGB")
    bitmap.save(out_path, quality=95)
    page.close()
    pdf.close()


def create_placeholder(path: Path, title: str, lines: list[str], width=1800, height=720) -> None:
    bg = (248, 249, 250)
    border = (110, 122, 140)
    text = (34, 45, 60)
    accent = (36, 98, 150)
    img = Image.new("RGB", (width, height), bg)
    draw = ImageDraw.Draw(img)
    try:
        font_title = ImageFont.truetype(FONT_CN_UI_BOLD, 54)
        font_body = ImageFont.truetype(FONT_CN_UI, 38)
    except OSError:
        font_title = ImageFont.load_default()
        font_body = ImageFont.load_default()
    draw.rounded_rectangle((20, 20, width - 20, height - 20), radius=24, outline=border, width=4)
    draw.rectangle((20, 20, width - 20, 128), fill=(232, 238, 245), outline=border, width=0)
    draw.text((70, 54), title, fill=accent, font=font_title)
    y = 190
    for line in lines:
        draw.text((95, y), line, fill=text, font=font_body)
        y += 68
    draw.text((95, height - 115), "后续获得实物图后，直接替换本占位图并保留图题。", fill=(120, 56, 56), font=font_body)
    img.save(path)


def center_text(draw, box, text, font, fill):
    left, top, right, bottom = box
    lines = text.split("\n")
    line_boxes = [draw.textbbox((0, 0), line, font=font) for line in lines]
    line_heights = [b[3] - b[1] for b in line_boxes]
    total_h = sum(line_heights) + (len(lines) - 1) * 14
    y = top + ((bottom - top) - total_h) / 2
    for line, b, h in zip(lines, line_boxes, line_heights):
        w = b[2] - b[0]
        x = left + ((right - left) - w) / 2
        draw.text((x, y), line, fill=fill, font=font)
        y += h + 14


def arrow(draw, start, end, fill=(58, 73, 92), width=7):
    x1, y1 = start
    x2, y2 = end
    draw.line((x1, y1, x2, y2), fill=fill, width=width)
    ang = math.atan2(y2 - y1, x2 - x1)
    length = 28
    spread = math.radians(28)
    pts = [
        (x2, y2),
        (x2 - length * math.cos(ang - spread), y2 - length * math.sin(ang - spread)),
        (x2 - length * math.cos(ang + spread), y2 - length * math.sin(ang + spread)),
    ]
    draw.polygon(pts, fill=fill)


def create_readable_flowchart(path: Path) -> None:
    width, height = 2100, 2750
    img = Image.new("RGB", (width, height), (250, 252, 255))
    draw = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype(FONT_CN_UI_BOLD, 72)
        node_font = ImageFont.truetype(FONT_CN_UI, 54)
        small_font = ImageFont.truetype(FONT_CN_UI, 42)
    except OSError:
        title_font = ImageFont.load_default()
        node_font = ImageFont.load_default()
        small_font = ImageFont.load_default()

    text_color = (22, 33, 49)
    stroke = (83, 104, 130)
    blue = (227, 241, 255)
    green = (227, 250, 238)
    purple = (242, 235, 255)
    yellow = (255, 246, 218)
    orange = (255, 237, 225)

    center_text(draw, (0, 52, width, 152), "上位机程序流程图", title_font, text_color)
    center_text(draw, (0, 142, width, 205), "主流程清晰展示，智能补扫作为旁路闭环", small_font, (83, 94, 112))

    x = 525
    w = 1050
    h = 150
    y0 = 255
    gap = 118
    nodes = [
        ("选择串口 / 设置扫频参数\n或导入文本数据", blue),
        ("发送 SWEEP 命令\n读取测量帧", blue),
        ("解析三数组\nomega / Magnitude / Phase", purple),
        ("数据校验与质量评估\nRMS、DC、削顶、有效采集", yellow),
        ("传统频响分析\n类型、阶次、关键频率", green),
        ("AI辅助诊断\n特征、模板拟合、知识库规则", green),
        ("汇总候选与证据链\n置信度、质量扣分、建议", purple),
    ]
    boxes = []
    for i, (label, fill) in enumerate(nodes):
        y = y0 + i * (h + gap)
        box = (x, y, x + w, y + h)
        boxes.append(box)
        draw.rounded_rectangle(box, radius=28, fill=fill, outline=stroke, width=5)
        center_text(draw, box, label, node_font, text_color)
        if i:
            prev = boxes[i - 1]
            arrow(draw, ((prev[0] + prev[2]) / 2, prev[3] + 6), ((box[0] + box[2]) / 2, box[1] - 6))

    # Decision diamond.
    y_dec = boxes[-1][3] + 130
    cx = width // 2
    diamond = [(cx, y_dec), (cx + 310, y_dec + 145), (cx, y_dec + 290), (cx - 310, y_dec + 145)]
    draw.polygon(diamond, fill=(255, 249, 234), outline=(221, 142, 35))
    draw.line(diamond + [diamond[0]], fill=(221, 142, 35), width=6)
    center_text(draw, (cx - 240, y_dec + 70, cx + 240, y_dec + 220), "需要补扫？", node_font, (128, 66, 35))
    arrow(draw, (cx, boxes[-1][3] + 6), (cx, y_dec - 6))

    # No branch.
    final_box = (x, y_dec + 410, x + w, y_dec + 560)
    draw.rounded_rectangle(final_box, radius=28, fill=blue, outline=(45, 118, 255), width=5)
    center_text(draw, final_box, "输出报告与文件\nBode / Nyquist、CSV、图像、诊断报告", node_font, text_color)
    arrow(draw, (cx - 85, y_dec + 290), (cx - 85, final_box[1] - 6))
    center_text(draw, (cx - 180, y_dec + 300, cx - 95, y_dec + 360), "否", small_font, (45, 118, 255))

    # Yes branch and loop.
    side_x = 1370
    side_w = 610
    step1 = (side_x, y_dec + 80, side_x + side_w, y_dec + 225)
    step2 = (side_x, y_dec + 335, side_x + side_w, y_dec + 480)
    for box, label in [
        (step1, "生成补扫计划\n1-3条 SWEEP"),
        (step2, "采集补扫数据\n合并并重新分析"),
    ]:
        draw.rounded_rectangle(box, radius=26, fill=orange, outline=(237, 99, 32), width=5)
        center_text(draw, box, label, small_font, text_color)
    arrow(draw, (cx + 310, y_dec + 145), (step1[0] - 8, (step1[1] + step1[3]) / 2))
    center_text(draw, (cx + 330, y_dec + 92, cx + 420, y_dec + 150), "是", small_font, (237, 99, 32))
    arrow(draw, ((step1[0] + step1[2]) / 2, step1[3] + 4), ((step2[0] + step2[2]) / 2, step2[1] - 6))
    draw.line((step2[0] - 12, (step2[1] + step2[3]) / 2, 168, (step2[1] + step2[3]) / 2, 168, boxes[3][1] + 75, boxes[3][0] - 8, boxes[3][1] + 75), fill=(83, 104, 130), width=6)
    arrow(draw, (boxes[3][0] - 8, boxes[3][1] + 75), (boxes[3][0] + 2, boxes[3][1] + 75), width=6)
    center_text(draw, (180, boxes[3][1] + 10, 520, boxes[3][1] + 94), "补扫后回到质量评估", small_font, (83, 94, 112))

    img.save(path)


def copy_and_prepare_assets() -> dict[str, Path]:
    assets = {}
    assets["system_block"] = ASSETS / "system_block_original.png"
    create_readable_flowchart(ASSETS / "program_flow_readable.png")
    assets["program_flow"] = ASSETS / "program_flow_readable.png"
    render_pdf_page(INPUTS / "schematic.pdf", 0, ASSETS / "schematic_page1.png", scale=4.0)
    assets["schematic"] = ASSETS / "schematic_page1.png"
    render_pdf_page(INPUTS / "pcb.pdf", 3, ASSETS / "pcb_top_label.png", scale=5.0)
    assets["pcb"] = ASSETS / "pcb_top_label.png"
    # The response plot is available from the earlier draft DOCX extraction.
    src_plot = ROOT / "source_image5.png"
    if src_plot.exists():
        assets["response_plot"] = src_plot
    create_placeholder(
        ASSETS / "hardware_photo_placeholder.png",
        "硬件实物图待补",
        ["待补：系统正面照片", "待补：系统斜 45°整体照片", "待补：电路板与接线局部照片"],
    )
    assets["hardware_placeholder"] = ASSETS / "hardware_photo_placeholder.png"
    create_placeholder(
        ASSETS / "software_screen_placeholder.png",
        "软件界面截图待补",
        ["待补：串口连接与扫频参数区", "待补：Bode / Nyquist 曲线区", "待补：智能诊断报告区"],
    )
    assets["software_placeholder"] = ASSETS / "software_screen_placeholder.png"
    return assets


def char_count(text: str) -> int:
    cleaned = re.sub(r"\s+", "", text)
    return len(cleaned)


LIMITED_BLOCKS: dict[str, tuple[str, int]] = {}


def limited(name: str, text: str, limit: int) -> str:
    LIMITED_BLOCKS[name] = (text, limit)
    return text


def audit_limits() -> None:
    problems = []
    for name, (text, limit) in LIMITED_BLOCKS.items():
        count = char_count(text)
        print(f"{name}: {count}/{limit}")
        if count > limit:
            problems.append(f"{name} {count}>{limit}")
    if problems:
        raise RuntimeError("字数超限: " + "; ".join(problems))


def build_doc() -> None:
    ensure_dirs()
    assets = copy_and_prepare_assets()

    doc = Document(str(TEMPLATE))
    clear_document_body(doc)
    configure_doc(doc)

    title = "基于STM32G431的智能频响测量与AI辅助电路诊断系统"
    p = doc.add_paragraph(style="Heading 2")
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_cm=None)
    r = p.add_run(title)
    set_run_font(r, size_pt=18, bold=True)

    add_text_para(doc, "摘要\t", size=14, bold=True, first_line=False)
    abstract = limited(
        "摘要",
        "本作品面向模拟滤波器、运放小信号网络及控制对象的频率响应测试，设计了基于STM32G431的智能频响测量与AI辅助电路诊断系统。系统由STM32测量端、可更换被测电路和Python上位机构成：PA4输出DAC扫频正弦激励，PA0采集输入参考，PA1采集输出响应，TIM2触发DAC DMA，TIM1同步触发双ADC采样；固件通过FFT提取幅值比和相位差，并以omega、Magnitude_data、Phase_data_rad三数组及诊断数组输出。上位机完成串口命令、文本导入、Bode/Nyquist绘图、传统频响分析、AI辅助诊断、主动补扫和报告导出。AI部分采用可解释规则与模板拟合结合的方法，从低/高频增益、斜率、-3 dB交点、相位跨度、RMS、DC偏置、削顶点等特征判断低通、高通、带通、带阻类型及阶次，并给出候选置信度、故障证据链和下一步补扫建议。系统已覆盖一阶至三阶低/高通及带通、带阻等典型对象，适合教学实验、嵌入式测量和电路调试场景。",
        800,
    )
    add_text_para(doc, abstract)

    add_section_heading(doc, "第一部分", "作品概述")

    add_subheading(doc, "功能与特性")
    add_text_para(
        doc,
        limited(
            "功能与特性",
            "系统支持用户在上位机设置起止频率、步进频率和激励幅值，一键发送SWEEP命令完成自动扫频。STM32端同步完成正弦激励、输入/输出采样、幅相计算和串口回传；上位机解析测量帧后显示Bode图、Nyquist图和诊断报告。作品的特性是软硬件链路闭环、协议文本可读、测量质量可追踪、AI结论可解释：当曲线异常时，系统不只给出类型标签，还会提示削顶、偏置、RMS过低、扫频范围不足或接线疑似反向等证据，帮助使用者快速定位问题。",
            400,
        )
    )
    add_figure(doc, assets["system_block"], "图1  系统测量链路与模块关系", width_cm=PAGE_USABLE_WIDTH_CM)

    add_subheading(doc, "应用领域")
    add_text_para(
        doc,
        limited(
            "应用领域",
            "本作品可用于低通、高通、带通、带阻等模拟滤波器的频响测试，也可扩展到运放小信号网络、控制对象开环测试和电子实验教学。传统手动扫频需要反复调节信号源、示波器并记录数据，容易漏测关键频段；本系统把扫频、采样、绘图、识别、诊断和补扫放入同一流程，适合课程实验、竞赛调试、PCB滤波模块验证和小型仪器开发。由于通信协议保持普通文本格式，测量结果还能直接用于离线复核、报告归档和二次算法验证。",
            400,
        )
    )

    add_subheading(doc, "主要技术特点")
    add_text_para(
        doc,
        limited(
            "主要技术特点",
            "（1）STM32G431片上DAC配合TIM2和DMA输出可调频率、幅值的正弦激励。（2）ADC1/ADC2由TIM1同步触发，PA0作为输入参考，PA1作为输出响应，幅相计算方向与传递函数定义一致。（3）串口协议保留omega、Magnitude_data、Phase_data_rad三数组，兼容文本导入和实时采集。（4）Python上位机集中完成绘图、分析、诊断、补扫和导出。（5）AI诊断采用特征提取、模板拟合、知识库规则和主动补扫组合，结论可复核。",
            400,
        )
    )

    add_subheading(doc, "主要性能指标")
    add_text_para(
        doc,
        limited(
            "主要性能指标说明",
            "系统以100 Hz-100 kHz典型扫频范围完成验证，串口115200 bps，ADC/DAC均为12位，单帧最多约200个频点，支持一阶至三阶低/高通及带通、带阻识别。",
            200,
        )
    )
    add_caption(doc, "表1  系统主要性能指标")
    add_table(
        doc,
        ["项目", "指标或实现情况"],
        [
            ["主控与转换", "STM32G431，12位DAC输出、12位双ADC采样"],
            ["默认扫频", "100 Hz-20 kHz，步进100 Hz，幅值1.2 Vpp"],
            ["验证范围", "10 Hz-100 kHz或100 Hz-100 kHz，按对象调整"],
            ["通信协议", "USART1，115200 bps，8N1，文本三数组"],
            ["算法输出", "类型、阶次、关键频率、候选置信度、故障证据、补扫计划"],
        ],
        [4.0, 10.2],
    )

    add_subheading(doc, "主要创新点")
    add_text_para(
        doc,
        limited(
            "主要创新点",
            "（1）把频响测量、质量评估和电路诊断组织为一条可闭环执行的测量链路。（2）以输入参考/输出响应的真实采样计算传递函数，降低理想激励假设带来的误差。（3）用可解释AI规则输出候选排名、故障证据和补扫命令，而不是只给单一分类结果。（4）知识库覆盖削顶、偏置、RMS过低、扫频范围不足等常见调试问题。",
            200,
        )
    )

    add_subheading(doc, "设计流程")
    add_text_para(
        doc,
        limited(
            "设计流程",
            "设计先确定测量对象、硬件通道和串口协议，再完成DAC扫频、双ADC同步采样与FFT幅相提取；随后构建上位机界面、传统频响分析和AI诊断规则；最后用多类滤波器数据验证识别、故障提示和补扫闭环。",
            200,
        )
    )

    add_section_heading(doc, "第二部分", "系统组成及功能说明")
    add_text_para(doc, "阐述具体的设计细节（图文结合）", first_line=True)

    add_subheading(doc, "整体介绍")
    add_text_para(
        doc,
        "系统由STM32G431测量端、被测滤波电路和PC上位机三部分组成。上位机发送扫频参数，STM32输出DAC激励并同步采样输入参考和输出响应，被测电路连接在激励输出与响应采样之间。测量端把角频率、幅值比、相位差及诊断信息回传给上位机，上位机完成曲线显示、传统分析、AI诊断、补扫决策和数据保存。各模块之间输入输出关系明确，便于硬件检查、软件调试和实验复现。"
    )
    add_caption(doc, "表2  系统主要组成模块")
    add_table(
        doc,
        ["模块", "输入", "输出", "功能说明"],
        [
            ["STM32测量端", "SWEEP命令、扫频参数", "三数组测量帧", "产生正弦激励，完成同步采样和幅相提取"],
            ["被测电路", "DAC扫频激励", "滤波响应信号", "形成待测传递函数"],
            ["协议解析层", "串口文本或导入文本", "MeasurementFrame", "提取测量数组及诊断数组"],
            ["传统分析层", "频率、幅值、相位", "类型、阶次、关键频率", "完成常规滤波器识别"],
            ["AI诊断层", "频响特征、诊断数组", "候选、证据、补扫计划", "形成可解释诊断结论"],
        ],
        [3.0, 3.5, 3.6, 4.1],
    )

    add_subheading(doc, "硬件系统介绍")
    add_text_para(doc, "2.2.1 硬件整体介绍", first_line=True)
    add_text_para(
        doc,
        "硬件系统以STM32G431为核心。PA4配置为DAC1_OUT1，经过TIM2 TRGO触发和DMA循环输出扫频正弦；PA0配置为ADC1输入参考，PA1配置为ADC2输出响应，两路ADC均由TIM1 TRGO上升沿触发。实际测试时，STM32开发板、被测滤波板和上位机共地连接，被测电路输出应保持在ADC允许范围内；交流耦合或运放滤波网络需要设置约1.65 V中点偏置，避免削顶和偏置漂移影响幅相计算。"
    )
    add_caption(doc, "表3  硬件外设与信号映射")
    add_table(
        doc,
        ["外设/信号", "引脚或配置", "作用"],
        [
            ["DAC1 CH1", "PA4，TIM2 TRGO，DMA循环", "产生扫频正弦激励"],
            ["ADC1", "PA0，TIM1 TRGO", "采集输入参考信号"],
            ["ADC2", "PA1，TIM1 TRGO", "采集输出响应信号"],
            ["USART1", "PA9/PA10，115200 bps", "接收命令并回传测量帧"],
            ["被测电路板", "LM358、RC网络、端子与排针", "实现低/高通、带通、带阻对象"],
        ],
        [3.5, 5.2, 5.5],
    )

    add_text_para(doc, "2.2.2 机械设计介绍", first_line=True)
    add_text_para(
        doc,
        "本作品当前重点为测量电路和软件系统，暂未设计独立机械外壳。后续实物完成后，将补充系统正面、斜45°整体照片、PCB安装位置、接口标识和线缆固定方式，用于说明便携性和操作安全性。"
    )

    add_text_para(doc, "2.2.3 电路各模块介绍", first_line=True)
    add_text_para(
        doc,
        "被测电路板包含LM358运放、RC网络、偏置电路、输入输出端子和多组排针，可组合低通、高通、带通、带阻等典型对象。原理图中标注了Vin、VOUT、BP_IN/BP_OUT、BR_IN/BR_OUT等关键节点；PCB图中预留+5V、GND、DA1、AD1、AD2等外部连接，便于与STM32的DAC输出和双ADC输入对应。"
    )
    add_figure(doc, assets["schematic"], "图2  被测电路原理图（SCH截图）", width_cm=PAGE_USABLE_WIDTH_CM)
    add_figure(doc, assets["pcb"], "图3  被测电路PCB关键布局图", width_cm=10.5)

    add_subheading(doc, "软件系统介绍")
    add_text_para(doc, "2.3.1 软件整体介绍", first_line=True)
    add_text_para(
        doc,
        "软件采用模块化组织。main.py作为入口，app.py构建Tkinter界面、参数输入、曲线绘制、历史记录和导出功能；serial_protocol.py解析omega、Magnitude_data、Phase_data_rad三数组并生成SWEEP命令；serial_readers.py与smart_resweep_reader.py负责串口读取、自动扫频和智能补扫线程；filter_analysis.py完成传统频响分析；ai_diagnosis.py结合特征、模板拟合和知识库规则形成AI辅助诊断；diagnosis_knowledge_base.json维护可扩展故障规则。"
    )
    add_caption(doc, "表4  上位机软件模块划分")
    add_table(
        doc,
        ["文件或模块", "功能说明"],
        [
            ["main.py / app.py", "启动GUI并组织串口、参数、绘图、诊断、导入导出流程"],
            ["serial_protocol.py", "解析三数组与诊断数组，生成SWEEP命令文本"],
            ["filter_analysis.py", "识别滤波器类型、阶次、关键频率和Nyquist稳定性信息"],
            ["ai_diagnosis.py", "提取特征、拟合模板、候选排序、规则诊断、生成补扫计划"],
            ["smart_resweep_reader.py", "执行1-3条补扫SWEEP并合并数据重新分析"],
            ["诊断知识库JSON", "维护削顶、偏置、RMS、噪声底、扫频范围等诊断规则"],
        ],
        [4.2, 10.0],
    )

    add_text_para(doc, "2.3.2 软件各模块介绍", first_line=True)
    add_text_para(
        doc,
        "程序流程从串口选择或文本导入开始。若执行实时测量，上位机根据用户参数发送SWEEP命令；若执行离线分析，软件直接读取文本三数组。数据进入MeasurementFrame后先完成长度、单调性、有效点、RMS、DC偏置和削顶检查，再分别进入传统频响分析和AI辅助诊断。AI诊断输出候选排名、模型拟合、质量扣分、故障证据链和补扫建议；若需要补扫，软件继续发送新的SWEEP命令并把补扫数据合并后重新分析。"
    )
    doc.add_page_break()
    add_figure(doc, assets["program_flow"], "图4  上位机程序流程图（按页宽放大，字体清晰可读）", width_cm=PAGE_USABLE_WIDTH_CM)
    doc.add_page_break()

    add_caption(doc, "表5  AI辅助诊断核心流程")
    add_table(
        doc,
        ["步骤", "输入", "输出"],
        [
            ["特征提取", "幅频、相频、RMS、DC、削顶、有效采集次数", "低/高频增益、斜率、相位跨度、-3 dB点等"],
            ["类型识别", "频响形态特征", "低通、高通、带通、带阻或未知响应"],
            ["模板拟合", "频率、幅值、相位", "阶次候选、中心频率、带宽、Q值"],
            ["故障规则", "诊断数组和知识库", "削顶、偏置、RMS过低、扫频不足等证据"],
            ["主动补扫", "关键频率和不确定区间", "1-3条新的SWEEP命令"],
        ],
        [3.2, 5.8, 5.2],
    )

    add_section_heading(doc, "第三部分", "完成情况及性能参数")
    add_text_para(doc, "阐述最终实现的成果（图文结合，实物照片为主）", first_line=True)

    add_subheading(doc, "整体介绍")
    add_text_para(
        doc,
        "系统已完成STM32测量端、Python上位机、传统频响分析、AI辅助诊断和智能补扫流程联调。用户完成接线后在上位机选择串口，设置扫频范围和激励幅值，点击开始扫频分析即可得到Bode图、Nyquist图和诊断文本。硬件实物照片待后续补充，本文先预留正面、斜45°和局部接线图位置。"
    )
    add_figure(doc, assets["hardware_placeholder"], "图5  系统实物照片占位（后续替换为正面、斜45°和局部接线图）", width_cm=PAGE_USABLE_WIDTH_CM)

    add_subheading(doc, "工程成果")
    add_text_para(doc, "3.2.1 机械成果", first_line=True)
    add_text_para(doc, "当前暂无独立机械结构，后续可补充外壳、固定孔位、面板端子和标签设计。")
    add_text_para(doc, "3.2.2 电路成果", first_line=True)
    add_text_para(doc, "电路成果包括STM32测量端连接关系、被测滤波PCB原理图与PCB布局。板上关键器件包含三片LM358、RC元件、输入输出端子和多组排针，可配合不同节点完成典型滤波网络验证。")
    add_text_para(doc, "3.2.3 软件成果", first_line=True)
    add_text_para(doc, "软件成果包括白色工程仪器风格GUI、串口扫频、文本导入、数据校验、曲线显示、智能诊断、智能补扫、CSV保存、图像保存和HTML报告导出。界面实物截图待最终演示环境稳定后补充。")
    add_figure(doc, assets["software_placeholder"], "图6  软件界面截图占位（后续替换为运行界面截图）", width_cm=PAGE_USABLE_WIDTH_CM)

    add_subheading(doc, "特性成果")
    if "response_plot" in assets:
        add_figure(doc, assets["response_plot"], "图7  八类典型滤波系统幅频响应测试结果", width_cm=PAGE_USABLE_WIDTH_CM)
    add_caption(doc, "表6  典型实测系统数据摘要")
    add_table(
        doc,
        ["系统", "点数", "扫频范围", "特征频率", "削顶情况"],
        [
            ["一阶低通", "207", "100.009 Hz-99116.785 Hz", "1586.433 Hz", "PA0=0，PA1=0"],
            ["二阶低通", "207", "100.009 Hz-99116.785 Hz", "1020.238 Hz", "PA0=0，PA1=0"],
            ["三阶低通", "207", "100.009 Hz-99116.785 Hz", "810.525 Hz", "PA0=0，PA1=0"],
            ["一阶高通", "359", "10.000 Hz-100000.005 Hz", "1317.785 Hz", "PA0=0，PA1=0"],
            ["二阶高通", "359", "10.000 Hz-100000.005 Hz", "1895.789 Hz", "PA0=0，PA1=0"],
            ["三阶高通", "359", "10.000 Hz-100000.005 Hz", "2309.026 Hz", "PA0=0，PA1=0"],
            ["二阶带阻", "359", "10.000 Hz-100000.005 Hz", "1960.784 Hz", "PA0=0，PA1=0"],
            ["二阶带通", "359", "10.000 Hz-100000.005 Hz", "3054.243 Hz", "PA0=0，PA1=0"],
        ],
        [2.6, 1.8, 4.2, 3.0, 2.6],
    )
    add_text_para(
        doc,
        "从测试结果看，系统能够覆盖低通、高通、带通和带阻等典型响应。低通类电路在高频端幅值滚降，高通类电路在低频端幅值较低，带通类电路在中心频率附近形成峰值，带阻类电路在阻带中心形成凹陷。AI诊断进一步结合候选置信度、质量扣分和故障证据，给出可操作的复测建议。"
    )

    add_section_heading(doc, "第四部分", "总结")

    add_subheading(doc, "可扩展之处")
    add_text_para(
        doc,
        limited(
            "可扩展之处",
            "系统后续可从硬件保护、测量范围和报告自动化继续完善。硬件方面，可在ADC输入端加入缓冲、限幅和可调偏置，提高对不同输出阻抗和不同幅值网络的适应能力；测量方面，可增加自动幅值控制，使阻带保持足够信噪比并避免通带削顶；软件方面，可把诊断报告自动整理成Word或PDF实验记录，并归档曲线、候选排名、故障证据和补扫结果。",
            300,
        )
    )

    add_subheading(doc, "心得体会")
    add_text_para(
        doc,
        limited(
            "心得体会",
            "本工程的关键体会是，频响测量系统不能只被看作画曲线的软件，也不能只被看作单片机采样程序，而应作为激励、采样、协议、分析、诊断和再测试构成的完整测量链路。当曲线出现异常时，问题可能来自接线方向、输入参考过低、输出削顶、ADC偏置、阻带噪声底、扫频范围不足或模型假设不匹配。如果系统只输出一张图，使用者仍需要凭经验判断；如果系统能够把异常拆成可观测特征，并用模板拟合和故障规则给出证据链，调试路径就会清楚许多。AI部分的设计也体现了同样思路：低通、高通、带通和带阻响应具有明确的电路规律，低/高频增益、幅频斜率、-3 dB交点、相位跨度、Nyquist曲线和RMS/DC诊断值都可以成为判断依据。因此，本系统没有把AI理解为不可解释的分类器，而是把工程师观察曲线、比较模板、排除测量异常、决定补扫区间的过程组织成程序流程。这样的AI结果可以被复核、被质疑、被重新测量验证，也更符合嵌入式测量仪器对可靠性和可解释性的要求。",
            1000,
        )
    )

    add_section_heading(doc, "第五部分", "参考文献")
    for ref in [
        "[1] STMicroelectronics. STM32G431xx Reference Manual and Datasheet.",
        "[2] Arm. CMSIS-DSP Software Library User Guide.",
        "[3] Alan V. Oppenheim, Ronald W. Schafer. Discrete-Time Signal Processing.",
        "[4] Katsuhiko Ogata. Modern Control Engineering.",
        "[5] Python Software Foundation. Python Documentation.",
        "[6] NumPy, SciPy and Matplotlib Documentation.",
    ]:
        add_text_para(doc, ref, first_line=False)

    audit_limits()
    doc.save(str(DOCX_OUT))
    print(DOCX_OUT)


if __name__ == "__main__":
    build_doc()
